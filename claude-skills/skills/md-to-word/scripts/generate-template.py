#!/usr/bin/env python3
"""
Generate a branded reference.docx for pandoc.

Reads all brand tokens dynamically from .bsg/DESIGN.md in the repo root.
Falls back to sensible neutral defaults when the file is absent or a token
is missing.  The template is project-agnostic: drop it in any BSG repo that
has a .bsg/DESIGN.md and it will pick up that project's identity.

Usage:
    python3 generate-template.py [output_path] [--logo path/to/logo.png]
    Default output: brand/templates/reference.docx
    Logo search order:
        1. --logo argument
        2. vitrine/public/<first *.png that looks like a logo>
        3. public/<logo*.png>
        4. brand/assets/logo.png
        5. No logo — company name used as text fallback
"""
import re
import sys
import subprocess
import tempfile
from pathlib import Path
from typing import Optional


def ensure_python_docx():
    try:
        import docx  # noqa: F401
    except ImportError:
        print("Installing python-docx…")
        subprocess.check_call([sys.executable, "-m", "pip", "install",
                               "python-docx", "-q"])


ensure_python_docx()

from docx import Document                         # noqa: E402
from docx.shared import Pt, Cm, Inches, RGBColor  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH     # noqa: E402
from docx.oxml.ns import qn                       # noqa: E402
from docx.oxml import OxmlElement                 # noqa: E402


# ── DESIGN.md parser ──────────────────────────────────────────────────────────

# Fallback tokens — neutral professional palette
_DEFAULTS = {
    "accent":        "#2563eb",
    "accent-dark":   "#1e3a8a",
    "accent-light":  "#3b82f6",
    "accent-violet": "#7c3aed",
    "foreground":    "#0f172a",
    "muted":         "#64748b",
    "border":        "#e2e8f0",
    "surface":       "#f1f5f9",
    "card":          "#ffffff",
    "font":          "Calibri",
    "font-mono":     "Courier New",
    "company":       "",
}

_COLOR_RE = re.compile(r"\|\s*`([^`]+)`\s*\|\s*`(#[0-9a-fA-F]{6})`")
_FONT_RE  = re.compile(r"\*\*([A-Za-z][A-Za-z0-9 ]+)\*\*.*(?:principal|primary|main)", re.I)
_H1_RE    = re.compile(r"^#\s+(.+?)(?:\s+—.*)?$", re.MULTILINE)


def parse_design_md(repo_root: Path) -> dict:
    """Extract brand tokens from .bsg/DESIGN.md.  Returns merged dict with defaults."""
    tokens = dict(_DEFAULTS)
    path = repo_root / ".bsg" / "DESIGN.md"
    if not path.exists():
        return tokens

    text = path.read_text(encoding="utf-8")

    # Colour tokens: | `token-name` | `#rrggbb` |
    for m in _COLOR_RE.finditer(text):
        tokens[m.group(1)] = m.group(2)

    # Primary font
    fm = _FONT_RE.search(text)
    if fm:
        tokens["font"] = fm.group(1).strip()

    # Company name: first H1, strip "— subtitle" suffix
    h1 = _H1_RE.search(text)
    if h1:
        tokens["company"] = h1.group(1).strip()

    return tokens


def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _rgb_str(hex_color: str) -> str:
    return hex_color.lstrip("#").upper()


# ── Logo discovery ────────────────────────────────────────────────────────────

_LOGO_GLOBS = [
    "vitrine/public/the-shift_dot_ai_petit.png",
    "vitrine/public/logo-*.png",
    "vitrine/public/*logo*.png",
    "public/logo*.png",
    "public/*logo*.png",
    "brand/assets/logo.png",
    "assets/logo.png",
]


def find_logo(repo_root: Path, hint: Optional[str] = None) -> Optional[str]:
    if hint:
        p = Path(hint)
        return str(p) if p.exists() else None
    for pattern in _LOGO_GLOBS:
        matches = sorted(repo_root.glob(pattern))
        if matches:
            return str(matches[0])
    return None


# ── XML helpers ───────────────────────────────────────────────────────────────

def _get_or_add(parent, tag: str):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el


# ── Style patching ────────────────────────────────────────────────────────────

def patch_style(doc, name, *, pt=None, bold=None, italic=None, color=None,
                font=None, before=None, after=None, indent_cm=None):
    try:
        s = doc.styles[name]
    except KeyError:
        return
    f = s.font
    if font  is not None: f.name       = font
    if pt    is not None: f.size       = Pt(pt)
    if bold  is not None: f.bold       = bold
    if italic is not None: f.italic    = italic
    if color is not None: f.color.rgb  = color
    if not hasattr(s, "paragraph_format"):
        return
    pf = s.paragraph_format
    if before    is not None: pf.space_before = Pt(before)
    if after     is not None: pf.space_after  = Pt(after)
    if indent_cm is not None: pf.left_indent  = Cm(indent_cm)


def set_line_spacing(doc, name, multiple=1.3):
    try:
        s = doc.styles[name]
    except KeyError:
        return
    pPr = s.element.get_or_add_pPr()
    sp = _get_or_add(pPr, "w:spacing")
    sp.set(qn("w:line"),     str(int(multiple * 240)))
    sp.set(qn("w:lineRule"), "auto")


# ── Table helpers ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _rgb_str(hex_color))
    tcPr.append(shd)


def set_cell_borders(cell, *, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    borders = OxmlElement("w:tcBorders")

    def _side(tag, spec):
        el = OxmlElement(f"w:{tag}")
        if spec is None:
            el.set(qn("w:val"), "none")
        else:
            color_hex, size = spec
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    str(size))
            el.set(qn("w:color"), color_hex)
            el.set(qn("w:space"), "0")
        borders.append(el)

    _side("top",    top)
    _side("bottom", bottom)
    _side("left",   left)
    _side("right",  right)
    tcPr.append(borders)


def set_cell_margins(cell, top=80, bottom=80, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"),    str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_no_spacing(tbl):
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblCellSpacing")):
        tblPr.remove(old)
    sp = OxmlElement("w:tblCellSpacing")
    sp.set(qn("w:w"),    "0")
    sp.set(qn("w:type"), "dxa")
    tblPr.append(sp)


def style_sample_table(tbl, tokens: dict):
    """Style the seed table in the reference.docx template."""
    set_table_no_spacing(tbl)
    accent  = tokens["accent"]
    border  = tokens["border"]
    surface = tokens["surface"]
    NONE  = None
    THIN  = (_rgb_str(border), 4)
    THICK = (_rgb_str(accent), 8)

    for r_idx, row in enumerate(tbl.rows):
        is_header   = r_idx == 0
        is_even     = r_idx % 2 == 0 and not is_header
        is_last_row = r_idx == len(tbl.rows) - 1
        bg         = accent   if is_header else (surface if is_even else "#ffffff")
        text_color = "#ffffff" if is_header else tokens["foreground"]

        for cell in row.cells:
            set_cell_bg(cell, bg)
            if is_header:
                set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                set_cell_borders(cell, top=THICK, bottom=THICK, left=THIN, right=THIN)
            else:
                set_cell_margins(cell, top=80, bottom=80, left=180, right=180)
                set_cell_borders(cell, top=NONE,
                                 bottom=THICK if is_last_row else THIN,
                                 left=THIN, right=THIN)
            for para in cell.paragraphs:
                para.clear()
                run = para.add_run("En-tête" if is_header else "Donnée")
                run.font.name      = tokens["font"]
                run.font.size      = Pt(10)
                run.font.bold      = is_header
                run.font.color.rgb = _rgb(text_color)


# ── Header / Footer ───────────────────────────────────────────────────────────

def add_header_footer(doc, tokens: dict, logo_path: Optional[str] = None):
    section = doc.sections[0]
    accent      = tokens["accent"]
    border      = tokens["border"]
    muted       = _rgb(tokens["muted"])
    accent_dark = _rgb(tokens["accent-dark"])
    font        = tokens["font"]
    company     = tokens.get("company", "")

    # ── Header ────────────────────────────────────────────────────────────────
    hdr = section.header
    p   = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    p.clear()

    if logo_path and Path(logo_path).exists():
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        try:
            run.add_picture(logo_path, height=Inches(0.35))
        except Exception:
            logo_path = None   # fall through to text fallback

    if not (logo_path and Path(logo_path).exists()):
        # Text fallback: company name or empty placeholder
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(company or " ")
        run.font.name      = font
        run.font.size      = Pt(9)
        run.font.bold      = bool(company)
        run.font.color.rgb = accent_dark

    # Accent bottom border on header paragraph
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), _rgb_str(accent))
    pBdr.append(bot)
    pPr.append(pBdr)

    # ── Footer ────────────────────────────────────────────────────────────────
    ftr = section.footer
    p2  = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    p2.clear()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pPr2  = p2._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    top_b = OxmlElement("w:top")
    top_b.set(qn("w:val"),   "single")
    top_b.set(qn("w:sz"),    "4")
    top_b.set(qn("w:space"), "4")
    top_b.set(qn("w:color"), _rgb_str(border))
    pBdr2.append(top_b)
    pPr2.append(pBdr2)

    run2 = p2.add_run()
    run2.font.name      = font
    run2.font.size      = Pt(9)
    run2.font.color.rgb = muted
    r = run2._r
    for tag in ("begin", "instr", "end"):
        if tag == "instr":
            el = OxmlElement("w:instrText")
            el.text = " PAGE "
            r.append(el)
        else:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"),
                   "separate" if tag == "end" else tag)
            r.append(fc)


# ── Pandoc discovery ──────────────────────────────────────────────────────────

def _find_pandoc() -> str:
    import shutil
    p = shutil.which("pandoc")
    if p:
        return p
    for c in ("/opt/homebrew/bin/pandoc", "/usr/local/bin/pandoc"):
        if Path(c).is_file():
            return c
    return "pandoc"


# ── Main ──────────────────────────────────────────────────────────────────────

def generate(output_path: str,
             logo_path:   Optional[str] = None,
             repo_root:   Optional[str] = None):

    root   = Path(repo_root) if repo_root else Path.cwd()
    tokens = parse_design_md(root)
    print(f"  Tokens: font={tokens['font']!r}  "
          f"accent={tokens['accent']}  company={tokens['company']!r}")

    logo = find_logo(root, logo_path)
    if logo:
        print(f"  Logo: {logo}")
    else:
        print(f"  Logo: aucun trouvé — fallback texte ({tokens['company'] or 'vide'})")

    pandoc = _find_pandoc()
    result = subprocess.run(
        [pandoc, "--print-default-data-file", "reference.docx"],
        capture_output=True,
    )
    if result.returncode != 0 or not result.stdout:
        print("Error: pandoc introuvable. Installer : brew install pandoc",
              file=sys.stderr)
        sys.exit(1)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(result.stdout)
        tmp_path = tmp.name

    doc = Document(tmp_path)

    # ── Page layout ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Typography (Tabula-inspired spacing: generous, grid-consistent) ────
    font      = tokens["font"]
    font_mono = tokens.get("font-mono", "Courier New")
    fg        = _rgb(tokens["foreground"])
    accent_dk = _rgb(tokens["accent-dark"])
    accent    = _rgb(tokens["accent"])
    muted     = _rgb(tokens["muted"])
    violet    = _rgb(tokens["accent-violet"])

    patch_style(doc, "Normal",
                font=font, pt=11, color=fg, after=10)
    set_line_spacing(doc, "Normal", 1.3)

    patch_style(doc, "Title",
                font=font, pt=26, bold=True, color=accent_dk,
                before=0, after=20)
    patch_style(doc, "Subtitle",
                font=font, pt=14, color=muted, after=24)

    patch_style(doc, "Heading 1",
                font=font, pt=22, bold=True, color=accent_dk,
                before=28, after=12)
    patch_style(doc, "Heading 2",
                font=font, pt=16, bold=True, color=accent,
                before=20, after=10)
    patch_style(doc, "Heading 3",
                font=font, pt=13, bold=True, color=fg,
                before=16, after=6)
    patch_style(doc, "Heading 4",
                font=font, pt=11, bold=True, italic=True, color=muted,
                before=12, after=4)

    for name in ("Block Text", "Block Quotation"):
        patch_style(doc, name, font=font, pt=11, color=muted,
                    indent_cm=1.0, before=6, after=6)

    for name in ("Source Code", "Verbatim"):
        patch_style(doc, name, font=font_mono, pt=9.5, color=accent_dk)
    for name in ("Verbatim Char", "Source Code Char"):
        patch_style(doc, name, font=font_mono, pt=10, color=violet)

    # ── Header / Footer ────────────────────────────────────────────────────
    add_header_footer(doc, tokens, logo)

    # ── Seed table (so pandoc picks up table cell styles) ─────────────────
    try:
        tbl = doc.add_table(rows=3, cols=3, style="Table")
        style_sample_table(tbl, tokens)
    except Exception as e:
        print(f"  (table seed ignorée : {e})")

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    Path(tmp_path).unlink(missing_ok=True)
    print(f"✓ Template généré : {out}")


if __name__ == "__main__":
    args      = sys.argv[1:]
    dest      = "brand/templates/reference.docx"
    logo_hint = None
    repo_hint = None
    i = 0
    while i < len(args):
        if args[i] == "--logo" and i + 1 < len(args):
            logo_hint = args[i + 1]; i += 2
        elif args[i] == "--repo" and i + 1 < len(args):
            repo_hint = args[i + 1]; i += 2
        else:
            dest = args[i]; i += 1

    generate(dest, logo_hint, repo_hint)
