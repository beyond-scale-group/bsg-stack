#!/usr/bin/env python3
"""
Post-process a pandoc-generated .docx: full-width tables + brand colours.

Reads brand tokens from .bsg/DESIGN.md in the repo root (falls back to
neutral defaults if absent).  Project-agnostic — works with any BSG repo.

Usage: python3 post-process.py <file.docx> [--repo /path/to/repo/root] [--logo /path/to/logo.png]

Logo selection order:
  1. --logo PATH (or $MD_TO_WORD_LOGO env var)
  2. a logo whose filename matches the DESIGN.md company (e.g. "BSG Holding"
     → logo-bsg-holding.png), searched in common brand directories
  3. the generic discovery globs (_LOGO_GLOBS)
  4. company name rendered as text (placeholder)
"""
import os
import re
import sys
from pathlib import Path


def ensure_python_docx():
    try:
        import docx  # noqa: F401
    except ImportError:
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install",
                               "python-docx", "-q"])


ensure_python_docx()

from docx import Document                         # noqa: E402
from docx.shared import Pt, Inches, RGBColor      # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH     # noqa: E402
from docx.oxml.ns import qn                       # noqa: E402
from docx.oxml import OxmlElement                 # noqa: E402


# ── DESIGN.md parser (shared with generate-template.py) ──────────────────────

_DEFAULTS = {
    "accent":        "#2563eb",
    "accent-dark":   "#1e3a8a",
    "foreground":    "#0f172a",
    "muted":         "#64748b",
    "border":        "#e2e8f0",
    "surface":       "#f1f5f9",
    "font":          "Calibri",
    "company":       "",
}
_COLOR_RE = re.compile(r"\|\s*`([^`]+)`\s*\|\s*`(#[0-9a-fA-F]{6})`")
_FONT_RE  = re.compile(r"\*\*([A-Za-z][A-Za-z0-9 ]+)\*\*.*(?:principal|primary|main)", re.I)
_H1_RE    = re.compile(r"^#\s+(.+?)(?:\s+—.*)?$", re.MULTILINE)


def parse_design_md(repo_root: Path) -> dict:
    tokens = dict(_DEFAULTS)
    path   = repo_root / ".bsg" / "DESIGN.md"
    if not path.exists():
        return tokens
    text = path.read_text(encoding="utf-8")
    for m in _COLOR_RE.finditer(text):
        tokens[m.group(1)] = m.group(2)
    fm = _FONT_RE.search(text)
    if fm:
        tokens["font"] = fm.group(1).strip()
    h1 = _H1_RE.search(text)
    if h1:
        tokens["company"] = h1.group(1).strip()
    return tokens


def _rgb_color(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _rgb_str(hex_color: str) -> str:
    return hex_color.lstrip("#").upper()


# ── Logo discovery ────────────────────────────────────────────────────────────

_LOGO_GLOBS = [
    "vitrine/public/the-shift_dot_ai_petit.png",
    "vitrine/public/logo-*.png",
    "vitrine/public/*logo*.png",
    "public/logo*.png",
    "public/*logo*.png",
    "brand/assets/logo.png",
    "assets/logo.png",
]


# Directories scanned when matching a logo to the DESIGN.md brand name.
_BRAND_DIRS = ["", "brand/assets", "brand", "public", "vitrine/public",
               "assets", "content/images", "content/strategy/images",
               "tools/document-generation/assets"]

# Generic words dropped before matching a logo filename, so the brand tokens
# drive the match regardless of the "Design System — Brand" / "Brand — …"
# heading convention.
_BRAND_STOPWORDS = {"design", "system", "the", "brand", "guide", "guidelines"}


def _brand_tokens(text: str):
    """"Design System — BSG Holding" -> ["bsg", "holding"]."""
    return [t for t in re.split(r"[^a-z0-9]+", text.lower())
            if t and t not in _BRAND_STOPWORDS]


def _design_h1(repo_root: Path) -> str:
    path = repo_root / ".bsg" / "DESIGN.md"
    if not path.exists():
        return ""
    m = re.search(r"^#\s+(.+?)\s*$", path.read_text(encoding="utf-8"),
                  re.MULTILINE)
    return m.group(1) if m else ""


def _brand_logo(repo_root: Path, company: str):
    """Find a *logo* whose filename contains every brand token."""
    tokens = _brand_tokens(_design_h1(repo_root) or company)
    if not tokens:
        return None
    candidates = []
    for d in _BRAND_DIRS:
        base = repo_root / d if d else repo_root
        if not base.is_dir():
            continue
        for p in base.glob("*.png"):
            name = p.name.lower()
            if "logo" in name and all(t in name for t in tokens):
                candidates.append(p)
    if candidates:
        # shortest filename = most specific brand logo; then alphabetical
        return str(sorted(candidates, key=lambda p: (len(p.name), p.name))[0])
    return None


def find_logo(repo_root: Path, override: str = "", company: str = ""):
    # 1. explicit override (CLI flag or env var)
    override = override or os.environ.get("MD_TO_WORD_LOGO", "")
    if override and Path(override).exists():
        return str(override)
    # 2. logo matching the DESIGN.md brand name
    by_brand = _brand_logo(repo_root, company)
    if by_brand:
        return by_brand
    # 3. generic discovery globs
    for pattern in _LOGO_GLOBS:
        matches = sorted(repo_root.glob(pattern))
        if matches:
            return str(matches[0])
    return None


# ── Cell helpers ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _rgb_str(hex_color))
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"),    str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_valign(cell, align="center"):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(old)
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), align)
    tcPr.append(va)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    borders = OxmlElement("w:tcBorders")

    def _side(tag, spec):
        el = OxmlElement(f"w:{tag}")
        if spec is None:
            el.set(qn("w:val"), "none")
        else:
            color_hex, size = spec
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    str(size))
            el.set(qn("w:color"), color_hex)
            el.set(qn("w:space"), "0")
        borders.append(el)

    _side("top",    top)
    _side("bottom", bottom)
    _side("left",   left)
    _side("right",  right)
    tcPr.append(borders)


# ── Table width ───────────────────────────────────────────────────────────────

def set_table_full_width(tbl):
    """Force table to 100 % text width, autofit columns to content."""
    tblEl = tbl._tbl
    tblPr = tblEl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tblEl.insert(0, tblPr)

    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)

    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tblPr.append(layout)

    for old in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(old)
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"),    "0")
    ind.set(qn("w:type"), "dxa")
    tblPr.append(ind)

    for old in tblPr.findall(qn("w:tblCellSpacing")):
        tblPr.remove(old)
    sp = OxmlElement("w:tblCellSpacing")
    sp.set(qn("w:w"),    "0")
    sp.set(qn("w:type"), "dxa")
    tblPr.append(sp)

    # Remove explicit cell widths — let autofit handle distribution
    seen = set()
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            if tc in seen:
                continue
            seen.add(tc)
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"),    "0")
            tcW.set(qn("w:type"), "auto")
            tcPr.append(tcW)


# ── Table brand styling ────────────────────────────────────────────────────────

def style_table_branded(tbl, tokens: dict):
    accent  = tokens["accent"]
    border  = tokens["border"]
    surface = tokens["surface"]
    fg      = tokens["foreground"]
    font    = tokens["font"]

    NONE  = None
    GRAY  = (_rgb_str(border), 4)
    THICK = (_rgb_str(accent), 8)

    n_rows   = len(tbl.rows)
    seen_tc  = set()

    for r_idx, row in enumerate(tbl.rows):
        is_header   = r_idx == 0
        is_even     = r_idx % 2 == 0 and not is_header
        is_last_row = r_idx == n_rows - 1
        bg         = accent   if is_header else (surface if is_even else "#ffffff")
        text_color = "#ffffff" if is_header else fg

        for cell in row.cells:
            tc = cell._tc
            if tc in seen_tc:
                continue
            seen_tc.add(tc)

            set_cell_bg(cell, bg)
            set_cell_valign(cell, "center")

            if is_header:
                set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
                set_cell_borders(cell, top=THICK, bottom=THICK, left=GRAY, right=GRAY)
            else:
                set_cell_margins(cell, top=100, bottom=100, left=200, right=200)
                set_cell_borders(cell, top=NONE,
                                 bottom=THICK if is_last_row else GRAY,
                                 left=GRAY, right=GRAY)

            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name      = font
                    run.font.size      = Pt(10)
                    run.font.color.rgb = _rgb_color(text_color)
                    if is_header:
                        run.font.bold = True
                if not para.runs:
                    r = para.add_run()
                    r.font.name      = font
                    r.font.size      = Pt(10)
                    r.font.color.rgb = _rgb_color(text_color)
                    if is_header:
                        r.font.bold = True


# ── Force TOC / fields update on open ────────────────────────────────────────

def enable_update_fields_on_open(doc: Document):
    """Set <w:updateFields w:val="true"/> in word/settings.xml.

    Without this flag, LibreOffice headless export (`soffice --convert-to pdf`)
    leaves the pandoc-generated TOC empty because it never recomputes fields.
    Word and LibreOffice both honour the flag and rebuild the TOC + page
    numbers when the document is opened — which is exactly what the PDF
    converter does behind the scenes.
    """
    settings_el = doc.settings.element
    for old in settings_el.findall(qn("w:updateFields")):
        settings_el.remove(old)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings_el.append(update)


# ── Logo before TOC ───────────────────────────────────────────────────────────

def insert_logo_before_toc(doc: Document, logo_path: str, company: str = ""):
    """Insert a centred logo (or company name) before the first body element."""
    body = doc.element.body

    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_para.paragraph_format.space_before = Pt(0)
    logo_para.paragraph_format.space_after  = Pt(32)

    inserted = False
    if logo_path and Path(logo_path).exists():
        run = logo_para.add_run()
        try:
            run.add_picture(logo_path, width=Inches(4.0))
            inserted = True
        except Exception as e:
            print(f"  (logo ignoré : {e})")

    if not inserted:
        if company:
            run = logo_para.add_run(company)
            run.font.size = Pt(18)
            run.font.bold = True
            inserted = True

    if not inserted:
        body.remove(logo_para._element)
        return

    logo_el     = logo_para._element
    body.remove(logo_el)
    first_child = next((el for el in body if el.tag != qn("w:sectPr")), None)
    if first_child is not None:
        body.insert(list(body).index(first_child), logo_el)
    else:
        body.insert(0, logo_el)

    label = logo_path if inserted and logo_path else company or "placeholder"
    print(f"✓ Couverture insérée avant le sommaire ({label})")


# ── Main ──────────────────────────────────────────────────────────────────────

def post_process(docx_path: str, repo_root: str = ".", logo_override: str = ""):
    root   = Path(repo_root)
    tokens = parse_design_md(root)
    logo   = find_logo(root, override=logo_override,
                       company=tokens.get("company", ""))

    doc = Document(docx_path)

    # 1. Logo / company name before TOC
    insert_logo_before_toc(doc, logo or "", tokens.get("company", ""))

    # 2. Full-width branded tables
    n = len(doc.tables)
    for tbl in doc.tables:
        set_table_full_width(tbl)
        style_table_branded(tbl, tokens)

    # 3. Tell readers (Word, LibreOffice) to refresh fields on open —
    #    required for the TOC to appear in headless PDF exports.
    enable_update_fields_on_open(doc)

    doc.save(docx_path)
    if n:
        print(f"✓ {n} tableau(x) mis en forme : {Path(docx_path).name}")


if __name__ == "__main__":
    args = sys.argv[1:]
    if not args:
        print("Usage: post-process.py <file.docx> [--repo /path/to/root] "
              "[--logo /path/to/logo.png]", file=sys.stderr)
        sys.exit(1)
    docx  = args[0]
    root  = "."
    logo  = ""
    if "--repo" in args:
        i    = args.index("--repo")
        root = args[i + 1]
    if "--logo" in args:
        i    = args.index("--logo")
        logo = args[i + 1]
    post_process(docx, root, logo_override=logo)
